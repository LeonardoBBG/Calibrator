"""DOCX, PDF and TXT ingestion with stable paragraph indexing."""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from .errors import DocumentParseError
from .models import Document, DocumentType, Paragraph

SUPPORTED_SUFFIXES = {".txt", ".docx", ".pdf"}
NUMBERED_RE = re.compile(
    r"^\s*(?:para(?:graph)?\s+)?(?:\((\d+(?:\.\d+)*)\)|(\d+(?:\.\d+)*))(?:[.)]|\s)+\s*(.+)$",
    re.IGNORECASE | re.DOTALL,
)


@dataclass(frozen=True)
class TextBlock:
    text: str
    is_heading: bool = False


def _normalise(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _looks_like_heading(text: str) -> bool:
    clean = _normalise(text)
    if not clean or NUMBERED_RE.match(clean) or len(clean) > 100:
        return False
    words = clean.split()
    return clean.endswith(":") or (len(words) <= 10 and clean.isupper())


def _plain_blocks(text: str) -> list[TextBlock]:
    blocks: list[TextBlock] = []
    pending: list[str] = []

    def flush() -> None:
        if pending:
            joined = _normalise(" ".join(pending))
            if joined:
                blocks.append(TextBlock(joined, _looks_like_heading(joined)))
            pending.clear()

    for raw_line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = _normalise(raw_line)
        if not line:
            flush()
        elif NUMBERED_RE.match(line):
            flush()
            blocks.append(TextBlock(line, False))
        else:
            pending.append(line)
    flush()
    return blocks


def _read_txt(path: Path) -> list[TextBlock]:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return _plain_blocks(data.decode(encoding))
        except UnicodeDecodeError:
            continue
    raise DocumentParseError(f"Unable to decode text file: {path}")


def _read_docx(path: Path) -> list[TextBlock]:
    try:
        document = DocxDocument(path)
    except Exception as exc:
        raise DocumentParseError(f"Unable to parse DOCX {path}: {exc}") from exc
    blocks: list[TextBlock] = []
    for paragraph in document.paragraphs:
        text = _normalise(paragraph.text)
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "").lower()
        blocks.append(TextBlock(text, style.startswith("heading") or _looks_like_heading(text)))
    return blocks


def _read_pdf(path: Path) -> list[TextBlock]:
    try:
        reader = PdfReader(str(path))
        page_text = []
        for page_number, page in enumerate(reader.pages, start=1):
            extracted = page.extract_text()
            if extracted is None:
                raise DocumentParseError(
                    f"PDF page {page_number} contains no extractable text; OCR is outside this MVP"
                )
            page_text.append(extracted)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Unable to parse PDF {path}: {exc}") from exc
    return _plain_blocks("\n\n".join(page_text))


def read_blocks(path: Path) -> list[TextBlock]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise DocumentParseError(
            f"Unsupported file format {suffix!r}; expected one of {sorted(SUPPORTED_SUFFIXES)}"
        )
    if not path.is_file():
        raise DocumentParseError(f"Input file does not exist: {path}")
    if suffix == ".txt":
        blocks = _read_txt(path)
    elif suffix == ".docx":
        blocks = _read_docx(path)
    else:
        blocks = _read_pdf(path)
    if not blocks:
        raise DocumentParseError(f"No usable text found in {path}")
    return blocks


def index_blocks(
    blocks: list[TextBlock], document_id: str, document_type: DocumentType
) -> list[Paragraph]:
    paragraphs: list[Paragraph] = []
    heading: str | None = None
    for block in blocks:
        raw = _normalise(block.text)
        match = NUMBERED_RE.match(raw)
        if block.is_heading and not match:
            heading = raw.rstrip(":")
            continue
        original_number: str | None = None
        text = raw
        if match:
            original_number = match.group(1) or match.group(2)
            text = _normalise(match.group(3))
        if not text:
            continue
        sequence = len(paragraphs) + 1
        paragraphs.append(
            Paragraph(
                paragraph_id=f"{document_type.value}-P{sequence:03d}",
                document_id=document_id,
                document_type=document_type,
                original_number=original_number,
                text=text,
                raw_text=raw,
                section_heading=heading,
                sequence=sequence,
            )
        )
    if not paragraphs:
        raise DocumentParseError("Document contained headings but no indexable paragraphs")
    return paragraphs


def ingest_document(
    path: Path, document_type: DocumentType, parser_version: str
) -> tuple[Document, list[Paragraph]]:
    file_bytes = path.read_bytes()
    digest = hashlib.sha256(file_bytes).hexdigest()
    parser_tag = hashlib.sha256(parser_version.encode("utf-8")).hexdigest()[:8].upper()
    document_id = f"DOC-{document_type.value}-{digest[:12].upper()}-{parser_tag}"
    document = Document(
        document_id=document_id,
        document_type=document_type,
        filename=path.name,
        sha256=digest,
        ingested_at=datetime.now(timezone.utc).isoformat(),
        parser_version=parser_version,
    )
    return document, index_blocks(read_blocks(path), document_id, document_type)
